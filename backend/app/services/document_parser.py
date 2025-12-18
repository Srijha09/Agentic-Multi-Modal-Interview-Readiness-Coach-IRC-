"""
Document parsing utilities for resumes and job descriptions.

Responsibilities:
- Load uploaded files from disk
- Extract raw text from PDF / DOCX / plain text
- Normalize whitespace
- Split into high‑level sections (Experience, Education, etc. for resumes;
  Responsibilities, Requirements, etc. for JDs) using simple heuristics
- Chunk sections into smaller pieces for downstream vectorization
"""
from __future__ import annotations

from pathlib import Path
from typing import List, Optional

from PyPDF2 import PdfReader  # type: ignore
from docx import Document as DocxDocument  # type: ignore

from app.core.config import settings
from app.schemas.document import (
    DocumentSection,
    DocumentChunk,
    ParsedDocument,
)


def _read_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(str(path))
    parts: List[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n".join(parts)


def _read_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    doc = DocxDocument(str(path))
    parts: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    return "\n".join(parts)


def _read_text(path: Path) -> str:
    """Read plain text file."""
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_text(path: Path) -> str:
    """
    Extract text from a file based on extension.

    Currently supports: .pdf, .docx, .txt
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix in {".docx", ".doc"}:
        return _read_docx(path)
    # Fallback to plain text
    return _read_text(path)


def normalize_whitespace(text: str) -> str:
    """Normalize whitespace: strip, collapse multiple spaces, normalize newlines."""
    # Normalize newlines first
    lines = [line.strip() for line in text.splitlines()]
    # Drop empty lines runs but keep paragraph breaks
    cleaned_lines: List[str] = []
    previous_blank = False
    for line in lines:
        if not line:
            if not previous_blank:
                cleaned_lines.append("")
            previous_blank = True
        else:
            cleaned_lines.append(line)
            previous_blank = False
    return "\n".join(cleaned_lines).strip()


def _guess_section_name(line: str, document_type: str) -> Optional[str]:
    """Heuristic to guess section name from a heading‑like line."""
    normalized = line.strip().lower().rstrip(":")
    if not normalized:
        return None

    resume_sections = {
        "experience",
        "work experience",
        "professional experience",
        "education",
        "projects",
        "skills",
        "summary",
        "objective",
        "certifications",
        "publications",
    }

    jd_sections = {
        "responsibilities",
        "requirements",
        "qualifications",
        "about the role",
        "about you",
        "nice to have",
        "benefits",
    }

    candidates = resume_sections if document_type == "resume" else jd_sections
    for section in candidates:
        if normalized == section or normalized.replace(" ", "") == section.replace(" ", ""):
            return section
    return None


def split_into_sections(text: str, document_type: str) -> List[DocumentSection]:
    """
    Split document text into high‑level sections using simple heading heuristics.

    Strategy:
    - Scan line by line
    - Treat lines in ALL CAPS or ending with ':' as potential headings
    - Map them to known section names (experience, education, etc.)
    - Group subsequent lines under that section
    """
    lines = text.splitlines()
    sections: List[DocumentSection] = []

    current_name: str = "body"
    current_lines: List[str] = []
    start_char = 0

    def flush_section(name: str, lines_for_section: List[str], start_offset: int):
        content = "\n".join(lines_for_section).strip()
        if not content:
            return start_offset
        end_offset = start_offset + len(content)
        sections.append(
            DocumentSection(
                name=name,
                text=content,
                start_char=start_offset,
                end_char=end_offset,
            )
        )
        return end_offset + 1  # plus newline

    for line in lines:
        stripped = line.strip()
        is_heading_candidate = (
            stripped.isupper()
            or stripped.endswith(":")
            or len(stripped.split()) <= 3
        )

        if is_heading_candidate:
            guessed = _guess_section_name(stripped, document_type=document_type)
            if guessed:
                # Flush previous section
                start_char = flush_section(current_name, current_lines, start_char)
                current_name = guessed
                current_lines = []
                continue

        current_lines.append(line)

    # Flush last section
    flush_section(current_name, current_lines, start_char)
    return sections


def chunk_section_text(
    section: DocumentSection,
    max_chars: int = 1000,
    overlap: int = 200,
) -> List[DocumentChunk]:
    """
    Chunk a section's text into overlapping windows.

    This is agnostic to tokens; a later phase can add token‑aware chunking.
    """
    text = section.text
    chunks: List[DocumentChunk] = []

    if not text:
        return chunks

    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        chunk_text = text[start:end]
        global_start = section.start_char + start
        global_end = section.start_char + end
        chunks.append(
            DocumentChunk(
                text=chunk_text,
                section_name=section.name,
                start_char=global_start,
                end_char=global_end,
            )
        )
        if end == len(text):
            break
        start = end - overlap
        if start < 0:
            start = 0

    return chunks


def parse_document(
    file_path: Path,
    document_type: str,
) -> ParsedDocument:
    """
    End‑to‑end parsing pipeline:
    - Extract raw text
    - Normalize whitespace
    - Split into sections
    - Chunk sections
    """
    raw_text = extract_text(file_path)
    normalized = normalize_whitespace(raw_text)
    sections = split_into_sections(normalized, document_type=document_type)

    all_chunks: List[DocumentChunk] = []
    for section in sections:
        all_chunks.extend(chunk_section_text(section))

    return ParsedDocument(
        document_type=document_type,
        file_name=file_path.name,
        text=normalized,
        sections=sections,
        chunks=all_chunks,
        char_count=len(normalized),
        section_count=len(sections),
        chunk_count=len(all_chunks),
    )


def get_upload_path(user_id: int, file_name: str) -> Path:
    """
    Compute and create the upload path for a given user + filename.
    """
    safe_name = file_name.replace("/", "_").replace("\\", "_")
    user_dir = settings.UPLOAD_DIR / f"user_{user_id}"
    user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir / safe_name

